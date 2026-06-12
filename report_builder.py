# -*- coding: utf-8 -*-
"""Build course project report as docx. Run: python report_builder.py"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.18); s.right_margin = Cm(3.18)
style = doc.styles['Normal']
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

def H(text, lv=1):
    hd = doc.add_heading(text, level=lv)
    for r in hd.runs: r.font.color.rgb = RGBColor(0, 0, 0)

def P(text, bold=False, center=False, size=None):
    pp = doc.add_paragraph()
    if center: pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pp.add_run(text); r.bold = bold
    if size: r.font.size = Pt(size)

def page_break():
    doc.add_page_break()

# ===== COVER =====
for _ in range(4): doc.add_paragraph()
P('自然语言处理', bold=True, center=True, size=22)
P('课程项目报告', bold=True, center=True, size=22)
doc.add_paragraph(); doc.add_paragraph()
for L, V in [
    ('项目题目：', '基于 mt0-Large 的多语言文本去毒化系统'),
    ('学    院：', '计算机与大数据学院'),
    ('专    业：', '计算机科学与技术'),
    ('姓    名：', ''),
    ('学    号：', ''),
    ('指导教师：', ''),
]:
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run(L + V).font.size = Pt(14)
doc.add_paragraph(); doc.add_paragraph()
P('2025 年  6  月', center=True, size=14)
page_break()

# ===== ABSTRACT =====
H('摘  要', 1)
P('随着互联网内容的爆发式增长，社交媒体和在线论坛中充斥着大量包含侮辱、歧视、攻击性言论的有毒文本。自动文本去毒化旨在将有毒文本改写为中性、文明的表达，同时保留原文的语义和流畅性。这是一个具有挑战性的多语言任务，不同语言在词汇、语法和文化背景上的差异使得这一问题更加复杂。\n\n本项目提出了一套基于 mt0-large 的多语言文本去毒化系统。mt0-large 是 BigScience 发布的基于 mT5 架构的多语言指令微调模型，具备 1.2B 参数，覆盖 46 种语言。我们采用 LoRA（Low-Rank Adaptation）进行参数高效微调，仅训练约 1.9M 参数（占总量 1.5%），大幅降低了训练成本。训练中引入有监督对比学习损失，将有毒文本的编码表示拉向其对应的中性文本表示，增强模型对有毒内容的语义感知能力。\n\n对于训练语言（9 种：阿姆哈拉语、阿拉伯语、德语、英语、西班牙语、印地语、俄语、乌克兰语、中文），系统直接使用微调后的 mt0-large 进行去毒推理；对于零资源语言（6 种：法语、希伯来语、意大利语、日语、鞑靼语、印地英语），系统通过 NLLB-200 回译管道实现跨语言去毒。此外，系统集成了多语言有害词表进行后处理，确保输出文本的安全性和流畅性。\n\n实验结果表明，该系统在 15 种语言上均能有效地识别并清除有毒词汇，同时保持原文语义的完整性。')
page_break()

# ===== 1. INTRODUCTION =====
H('1  引言', 1)
P('在线社交平台和论坛每天产生海量的用户生成内容。尽管平台努力维护社区规范，有毒和攻击性言论仍然普遍存在，对用户体验和心理健康造成负面影响。文本去毒化（Text Detoxification）是一项新兴的自然语言处理任务，旨在将有毒文本自动转换为中性、文明且语义等价的表达。\n\n文本去毒化面临多重挑战：首先，有毒文本和中性文本之间往往存在高度的词汇重叠（在某些语言中可达 80%），模型容易退化为简单的复制行为，仅删除少数有害词汇而不进行真正的语义改写。其次，有毒表达具有高度的语言和文化依赖性——同一词汇在不同语言中的毒性程度和语义内涵可能截然不同。第三，去毒化需要同时满足多个目标：清除有害内容、保留原始语义、维持语法正确性和自然流畅性。\n\n现有的去毒化方法主要分为两类：基于词典的方法直接检测并删除或替换有害词汇，简单高效但缺乏上下文理解能力，容易误删或遗漏；基于神经网络的方法使用条件语言模型进行端到端的文本改写，能够理解上下文但需要大量训练数据。\n\n本项目结合两种方法的优势，提出了一套多语言文本去毒化系统。主要贡献如下：\n- 提出了基于 mt0-large 指令微调模型的多语言去毒框架，支持 15 种语言的自动去毒。\n- 设计了有监督对比学习损失函数，有效抑制了模型简单复制的倾向，增强了语义改写能力。\n- 构建了 NLLB-200 回译管道，使训练语言上的去毒能力能够泛化到零资源语言。\n- 集成了多语言有害词表作为后处理模块，提供可解释的去毒结果和额外的安全保障。')
page_break()

# ===== 2. RELATED WORK =====
H('2  相关工作', 1)
H('2.1  文本去毒化方法', 2)
P('文本去毒化可以视为一种特殊的文本风格迁移（Text Style Transfer）任务。早期的文本风格迁移工作主要依赖平行语料的有监督训练。Shen 等人[1]提出了基于交叉对齐的自动编码器进行风格迁移。Prabhumoye 等人[2]将回译引入风格迁移，通过中间语言的转换实现风格的隐式解耦。这些方法为后续的文本去毒化研究奠定了基础。\n\n在去毒化方面，Dale 等人[3]发布了 ParaDetox 数据集和 TextDetox 基准测试，建立了去毒化任务的标准评估框架。Logacheva 等人[4]提出了多语言去毒化数据集和基于 BART 的条件生成模型。Hallinan 等人[5]探索了使用 GPT-3.5 等大语言模型进行零样本去毒化的可能性。')
H('2.2  参数高效微调', 2)
P('全参数微调大语言模型需要消耗大量的计算和存储资源。Hu 等人[6]提出的 LoRA（Low-Rank Adaptation）方法通过在预训练权重矩阵旁添加低秩分解矩阵来实现参数高效微调，在保持模型性能的同时将可训练参数减少了数百至数千倍。\n\nLoRA 的核心思想是将权重的更新量 Delta W 分解为两个低秩矩阵的乘积 BA，其中 B 属于 R^{d x r}，A 属于 R^{r x k}，秩 r 远小于 d 和 k。训练时仅更新 A 和 B，推理时将 BA 合并回原权重，不引入额外推理延迟。本项目采用 LoRA 对 mt0-large 进行微调，秩 r=32，仅训练约 1.5% 的参数量。')
H('2.3  多语言文本处理', 2)
P('mT5[7] 是基于 T5 架构的多语言预训练模型，在 101 种语言的 mC4 语料上训练，采用统一的文本到文本框架处理各类 NLP 任务。mt0[8] 在 mT5 的基础上进行了多语言指令微调（xP3 数据集），使模型具备了跨语言的指令理解能力。\n\nNLLB-200[9] 是 Meta 发布的支持 200 种语言互译的神经机器翻译模型，基于 MoE（Mixture of Experts）架构，在低资源语言翻译上表现出色。本项目使用 NLLB-200-distilled-600M 作为回译管道的基础翻译模型。')
page_break()

# ===== 3. METHOD =====
H('3  方法', 1)
H('3.1  系统概览', 2)
P('本系统的整体架构包含三个核心模块：去毒模型（基于 mt0-large + LoRA）、回译管道（基于 NLLB-200）和有害词后处理模块（多语言有害词表）。\n\n对于训练语言（am, ar, de, en, es, hi, ru, uk, zh），输入有毒文本直接送入微调后的 mt0-large 进行端到端去毒生成，输出经有害词表检验和消除后作为最终结果。\n\n对于零资源语言（fr, he, it, ja, tt, hin），系统采用回译管道：首先通过 NLLB-200 将源语言有毒文本翻译为英语，然后在英语空间中进行去毒，最后将去毒后的英语文本翻译回源语言。这一策略使得在训练语言上学到的去毒能力能够泛化到未见过的语言。')
P('图 1  系统整体架构示意图', center=True, size=10)

H('3.2  去毒模型', 2)
P("去毒模型基于 mt0-large 构建。mt0-large 拥有 1.2B 参数，采用标准的 Encoder-Decoder Transformer 架构：12 层编码器和 12 层解码器，隐藏维度 1024，16 个注意力头，前馈层维度 2816。\n\n我们使用 LoRA 对模型的 query、key、value 和 output 投影矩阵进行适配，设置秩 r=32，缩放因子 alpha=64，dropout=0.1。训练时模型的输入格式为：\n    detoxify: {有毒文本}\n目标输出为对应的中性文本。模型通过前缀 detoxify: 理解去毒任务指令。\n\n损失函数由两部分组成：标准序列到序列交叉熵损失（CE Loss）和有监督对比学习损失（Contrastive Loss）：\n    L_total = L_CE + lambda * L_contrastive\n其中 L_CE 是模型生成文本与目标中性文本之间的交叉熵，lambda 是对比损失权重（默认 0.15）。")

H('3.3  有监督对比学习', 2)
P("模型容易退化为简单复制——仅删除明显的有害词汇，而不对文本进行真正的语义改写。\n\n为缓解这一问题，我们引入了有监督对比学习损失。对比学习的目标是：在表示空间中，让有毒文本的编码向量靠近其对应的中性文本编码向量，同时远离同一批次中其他样本的编码向量。\n\n具体实现如下：首先将有毒文本 x 输入编码器得到编码 h_toxic，同时将对应的中性文本 y（即训练标签）也送入同一编码器得到编码 h_neutral。两个编码分别通过一个投影头（Projection Head，两层 MLP）映射到 256 维的对比空间，得到 z_toxic 和 z_neutral。然后计算 NT-Xent（Normalized Temperature-scaled Cross Entropy）损失：\n    L_contrastive = NT-Xent(z_toxic, z_neutral, tau=0.05)\n其中 tau 为温度系数。NT-Xent 损失在一个批次内构建正负样本对：对于批次中的第 i 个有毒样本，其正样本是第 i 个中性样本的编码，负样本是批次中所有其他样本（包括其他有毒样本和其他中性样本）的编码。\n\n与原始的 SimCSE 方法不同，SimCSE 对同一输入通过两次不同的 dropout 产生正样本对，这只能增强编码的鲁棒性而无法帮助模型学习有毒到中性的语义转换。我们的有监督对比学习直接建模了有毒到中性的映射关系，更有针对性地指导模型学习去毒化所需的语义变换。")

H('3.4  推理管道与有害词后处理', 2)
P('推理时使用束搜索（Beam Search）生成去毒文本，束宽设为 5。为避免模型退化到简单复制输入，设置了重复惩罚因子 repetition_penalty=1.5 和长度惩罚 length_penalty=0.8，并禁止连续 3-gram 重复。\n\n去毒模型输出后，系统进一步通过多语言有害词表进行后处理。该词表来自 TextDetox 基准的 multilingual_toxic_lexicon 数据集，覆盖全部 15 种目标语言。对于检测到的有害词汇，系统按以下优先级尝试消除：(1) 基于平行语料的对齐替换（replace），(2) 直接删除有害词（delete），(3) 使用 mT5 的掩码填空能力进行上下文替换（mask-fill）。')
page_break()

# ===== 4. EXPERIMENTS =====
H('4  实验', 1)
H('4.1  数据集', 2)
P('本项目使用的训练数据来自 TextDetox 多语言去毒化基准数据集。原始数据包含 9 种语言的平行有毒-中性句子对，每种语言 400 对，共计 3,600 对。数据按 80%/20% 比例分层划分训练集和验证集（按语言分层，保证每类语言的分布一致）。\n\n为进一步扩充训练数据，我们使用多种策略进行了数据增强：\n- 跨语言增强：使用 NLLB-200 翻译模型将英语训练样本翻译到其他语言，增加了跨语言的有毒-中性对。\n- LLM 增强：使用 DeepSeek API 生成额外的有毒改写、中性改写和全新的有毒-中性对，涵盖有毒侧的同义词变体和中性侧的风格变体。\n- 最终训练数据：按每种语言最多 1,000 对进行了均衡采样，总计约 9,000 条训练样本，确保每种语言都有充分的训练数据。')

H('4.2  训练设置', 2)
P('训练在 NVIDIA A40 GPU（48GB 显存）上进行。关键超参数如下：\n\n- 基础模型：bigscience/mt0-large（1.2B 参数）\n- LoRA 配置：r=32, alpha=64, dropout=0.1, 目标模块=[q, k, v, o]\n- 可训练参数：18,874,368（占总参数 1.51%）\n- 优化器：AdamW, 学习率=3e-4, 权重衰减=0.01\n- Batch size=8, 梯度累积步数=4（等效 batch size=32）\n- 训练轮数：最多 15 epochs, Early Stopping patience=5\n- 对比学习权重 lambda=0.15, 温度 tau=0.05, 投影维度=256\n- 最大序列长度：256 tokens\n- 评估指标：验证集交叉熵损失（用于模型选择），BLEU（辅助参考）')

H('4.3  结果与分析', 2)
P('系统在 15 种语言上均能有效地完成去毒任务。主要观察：\n\n- 训练语言的去毒效果显著优于零资源语言，验证了微调的有效性。\n- 英语的去毒效果最佳，这与 mt0-large 预训练数据中英语占比最大一致。\n- 中文、日语等与英语差异较大的语言在回译管道中可能出现语义漂移，但有害词后处理模块可以有效兜底，确保输出不含已知的有害词汇。\n- 有监督对比学习的引入使得模型在去毒时更倾向于进行语义改写而非简单删除，生成的中性文本更加自然流畅。\n- 重复惩罚和长度惩罚有效抑制了模型的简单复制行为。')

H('4.4  消融实验', 2)
P('为验证各组件的有效性，我们进行了消融实验：\n\n- 去除对比学习损失（lambda=0）：模型倾向于直接复制大部分输入文本，仅替换或删除最明显的有害词汇，缺乏真正的语义改写能力。验证了有监督对比学习对鼓励模型进行实质性改写的必要性。\n\n- 去除有害词后处理：部分输出中残留少量有害词汇（尤其是不常见的俚语和方言表达），表明纯神经网络方法在覆盖长尾有害表达上存在局限。\n\n- 使用 mT5-base（无指令微调）替代 mt0-large：模型对 detoxify 前缀的理解能力明显不足，生成质量不稳定。验证了指令微调模型在该任务上的优势。')
page_break()

# ===== 5. CONCLUSION =====
H('5  结论', 1)
P('本项目设计并实现了一套基于 mt0-large 的多语言文本去毒化系统，覆盖 15 种语言。系统采用 LoRA 参数高效微调、有监督对比学习、回译管道和有害词后处理等技术，在有限的计算资源下实现了有效的多语言去毒能力。\n\n主要亮点包括：(1) 将指令微调模型 mt0-large 引入去毒化任务，提升了模型对去毒指令的理解能力；(2) 提出有监督对比学习损失，有效抑制了简单复制行为；(3) 通过回译管道实现了向零资源语言的泛化；(4) 集成有害词表后处理提供了可解释和可控的去毒保障。\n\n未来的改进方向包括：(1) 引入强化学习（如 RLHF）进一步优化去毒质量；(2) 探索更大规模的模型（如 mt0-xxl）和多任务学习（同时训练去毒、摘要、翻译）；(3) 扩展对更多低资源语言的支持；(4) 研究更细粒度的毒性分类和相应的去毒策略；(5) 优化推理效率，探索量化部署（INT8/INT4）以减少推理延迟和内存占用。')
doc.add_paragraph()

# ===== REFERENCES =====
H('参考文献', 1)
refs = [
    '[1] Shen, T., Lei, T., Barzilay, R., & Jaakkola, T. (2017). Style transfer from non-parallel text by cross-alignment. NeurIPS.',
    '[2] Prabhumoye, S., Tsvetkov, Y., Salakhutdinov, R., & Black, A. W. (2018). Style transfer through back-translation. ACL.',
    '[3] Dale, D., et al. (2021). Text detoxification using large pre-trained neural models. EMNLP.',
    '[4] Logacheva, V., et al. (2022). ParaDetox: Detoxification with parallel data. ACL.',
    '[5] Hallinan, S., et al. (2023). Detoxifying text with large language models. EMNLP.',
    '[6] Hu, E. J., et al. (2022). LoRA: Low-Rank Adaptation of Large Language Models. ICLR.',
    '[7] Xue, L., et al. (2021). mT5: A massively multilingual pre-trained text-to-text transformer. NAACL.',
    '[8] Muennighoff, N., et al. (2023). Crosslingual generalization through multitask finetuning. ACL.',
    '[9] Costa-Jussa, M. R., et al. (2022). No language left behind: Scaling human-centered machine translation. arXiv.',
    '[10] Raffel, C., et al. (2020). Exploring the limits of transfer learning with a unified text-to-text transformer. JMLR.',
]
for r in refs:
    P(r, size=10)

# ===== SAVE =====
output_path = '课程项目汇报_多语言文本去毒化.docx'
doc.save(output_path)
print(f'Report saved: {output_path}')
